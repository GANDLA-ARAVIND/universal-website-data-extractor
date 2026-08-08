import io
from typing import Any, List, Optional, Tuple

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from src.application.services.exporters.base import BaseExporter
from src.db.models.crawl_job import CrawlJob
from src.db.models.page import ExtractedPage
from src.db.models.statistic import CrawlStatistic
from src.schemas.export import ExportFormat


from src.schemas.dataset import StandardCrawlDataset


def add_single_dataset_docx(doc: Document, dataset: StandardCrawlDataset) -> None:
    """Appends a complete 9-section single website extraction report to a Word document."""
    # Title Header
    title_p = doc.add_paragraph()
    title_run = title_p.add_run(f"Website Extraction Report: {dataset.website_info.domain}")
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    sub_p = doc.add_paragraph()
    sub_run = sub_p.add_run(f"Target Seed URL: {dataset.website_info.seed_url}")
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
    sub_run.font.bold = True

    # 1. Executive Summary
    h1 = doc.add_heading("1. Executive Summary", level=1)
    h1.style.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)

    table_exec = doc.add_table(rows=6, cols=2)
    table_exec.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_exec.autofit = False

    exec_rows = [
        ("Crawl Status", dataset.website_info.status),
        ("Total Pages Extracted", str(dataset.statistics.pages_crawled)),
        ("Failed Pages Count", str(dataset.statistics.failed_pages)),
        ("Total Images Discovered", str(dataset.statistics.total_images)),
        ("Total Links Discovered", str(dataset.statistics.total_links)),
        ("Crawl Execution Duration", f"{dataset.statistics.total_duration_sec}s"),
    ]
    for idx, (k, v) in enumerate(exec_rows):
        row = table_exec.rows[idx]
        row.cells[0].paragraphs[0].add_run(k).bold = True
        row.cells[1].paragraphs[0].add_run(v)
        row.cells[0].width = Inches(2.5)
        row.cells[1].width = Inches(4.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # 2. Website Overview
    h2 = doc.add_heading("2. Website Overview", level=1)
    h2.style.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)

    p_over = doc.add_paragraph()
    p_over.add_run("Website Title: ").bold = True
    p_over.add_run(dataset.summary.title or "N/A")

    p_desc = doc.add_paragraph()
    p_desc.add_run("Meta Description: ").bold = True
    p_desc.add_run(dataset.summary.meta_description or "N/A").italic = True

    p_meta = doc.add_paragraph()
    p_meta.add_run(f"Target Domain: {dataset.website_info.domain} | Max Depth: {dataset.metadata.max_depth} | Total Headings: {dataset.summary.total_headings_found} | Total Tables: {dataset.summary.total_tables_found}")

    if dataset.summary.main_sections:
        p_sec = doc.add_paragraph()
        p_sec.add_run("Main Discovered Sections: ").bold = True
        p_sec.add_run(", ".join(dataset.summary.main_sections))

    # 3. Website Structure
    h3 = doc.add_heading("3. Website Structure Hierarchy", level=1)
    h3.style.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
    for line in dataset.summary.structure_tree:
        p_tree = doc.add_paragraph(line)
        p_tree.style.font.name = "Consolas"
        p_tree.style.font.size = Pt(9.5)

    # 4. Discovered Page Index
    h4 = doc.add_heading("4. Discovered Page Index", level=1)
    h4.style.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)

    idx_rows = len(dataset.pages) + 1
    table_idx = doc.add_table(rows=idx_rows, cols=6)
    table_idx.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["#", "Title", "URL", "Status", "Depth", "Latency"]
    for c_idx, h_text in enumerate(headers):
        cell_p = table_idx.rows[0].cells[c_idx].paragraphs[0]
        run = cell_p.add_run(h_text)
        run.bold = True
        run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    for r_idx, p in enumerate(dataset.pages, 1):
        row_cells = table_idx.rows[r_idx].cells
        row_cells[0].paragraphs[0].add_run(str(r_idx))
        row_cells[1].paragraphs[0].add_run((p.title or "Untitled")[:30])
        row_cells[2].paragraphs[0].add_run(p.url[:40])
        row_cells[3].paragraphs[0].add_run(str(p.status_code))
        row_cells[4].paragraphs[0].add_run(str(p.depth))
        row_cells[5].paragraphs[0].add_run(f"{p.response_time_ms:.0f}ms")

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # 5. Detailed Page Information
    h5 = doc.add_heading("5. Detailed Page Information", level=1)
    h5.style.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)

    for idx, p in enumerate(dataset.pages, 1):
        doc.add_heading(f"Page {idx}: {p.title or 'Untitled Page'}", level=2)
        meta_p = doc.add_paragraph()
        meta_p.add_run("URL: ").bold = True
        meta_p.add_run(p.url)
        meta_p.add_run(f" | Status: {p.status_code} | Depth: {p.depth} | Latency: {p.response_time_ms:.1f}ms")

        if p.meta_description:
            dp = doc.add_paragraph()
            dp.add_run("Description: ").bold = True
            dp.add_run(p.meta_description).italic = True

        if p.headings:
            doc.add_heading("Headings", level=3)
            for level, txts in p.headings.items():
                for t in txts[:5]:
                    doc.add_paragraph(f"[{level.upper()}] {t}", style="List Bullet")

        if p.paragraphs:
            doc.add_heading("Content Paragraphs", level=3)
            for para in p.paragraphs[:3]:
                doc.add_paragraph(para)

    # 6. Errors & Warnings
    h6 = doc.add_heading("6. Execution Errors & Warnings", level=1)
    h6.style.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
    if dataset.errors or dataset.warnings:
        for err in dataset.errors:
            doc.add_paragraph(f"[ERROR] {err}", style="List Bullet")
        for warn in dataset.warnings:
            doc.add_paragraph(f"[WARNING] {warn}", style="List Bullet")
    else:
        doc.add_paragraph("No execution errors or warnings recorded.")

    # 7. Appendix
    h7 = doc.add_heading("7. Technical Appendix", level=1)
    h7.style.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
    doc.add_paragraph(f"Generated At: {dataset.download_metadata.get('generated_at', 'N/A')}")
    doc.add_paragraph(f"Exporter Version: {dataset.download_metadata.get('exporter_version', '3.0.0')}")
    doc.add_paragraph("Schema Standard: StandardCrawlDataset v3.0")


class DocxExporter(BaseExporter):
    """Strategy Exporter for professional Microsoft Word (.docx) document generation."""

    @property
    def format_type(self) -> ExportFormat:
        return ExportFormat.DOCX

    async def export(
        self,
        pages: List[ExtractedPage],
        job: CrawlJob,
        stats: Optional[CrawlStatistic] = None,
    ) -> Tuple[bytes, str, str]:
        dataset = StandardCrawlDataset.from_orm_models(pages, job, stats)
        doc = Document()

        for section in doc.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        add_single_dataset_docx(doc, dataset)

        buffer = io.BytesIO()
        doc.save(buffer)
        docx_bytes = buffer.getvalue()
        buffer.close()

        filename = self.generate_filename(job.seed_url, str(job.id), "docx")
        return (
            docx_bytes,
            filename,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    async def export_batch_dataset(
        self,
        batch_dataset: Any,
    ) -> Tuple[bytes, str, str]:
        """Generates Word document (.docx) payload for a BatchDataset."""
        doc = Document()
        for section in doc.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        title_p = doc.add_paragraph()
        run = title_p.add_run("Multi-Website Batch Crawl Executive Report")
        run.font.size = Pt(22)
        run.bold = True

        doc.add_paragraph(f"Batch ID: {batch_dataset.batch_metadata.batch_id}")
        doc.add_paragraph(f"Overall Status: {batch_dataset.batch_statistics.overall_status}")
        doc.add_paragraph(f"Total Target Websites: {batch_dataset.batch_statistics.total_websites} | Successful: {batch_dataset.batch_statistics.successful_websites} | Failed: {batch_dataset.batch_statistics.failed_websites}")
        doc.add_paragraph(f"Total Pages Extracted: {batch_dataset.batch_statistics.total_pages} | Duration: {batch_dataset.batch_statistics.total_duration_sec}s")

        doc.add_page_break()

        for site in batch_dataset.websites:
            if site.dataset:
                add_single_dataset_docx(doc, site.dataset)
                doc.add_page_break()

        buffer = io.BytesIO()
        doc.save(buffer)
        docx_bytes = buffer.getvalue()
        buffer.close()

        short_id = str(batch_dataset.batch_metadata.batch_id)[:8]
        filename = f"batch_export_{short_id}.docx"
        return (
            docx_bytes,
            filename,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )


